from flask import Flask, request, send_file, render_template, jsonify
from docx import Document
import os
from datetime import datetime, timedelta
from num2words import num2words
from docx.shared import Pt
from dateutil.relativedelta import relativedelta
import calendar
import logging
import re
import inflect
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_UNDERLINE

# call inflect
p = inflect.engine()

# At the top of your script
logging.basicConfig(level=logging.INFO,
                    format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

app = Flask(__name__)

# Helper function to parse dates
def parse_and_format_date(date_str):
    """Parse dates in multiple formats and return both datetime object and formatted string"""
    formats = ["%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y"]

    for fmt in formats:
        try:
            date_obj = datetime.strptime(date_str, fmt)
            return date_obj, date_obj.strftime("%d/%m/%Y")
        except ValueError:
            continue

    logger.error(f"Invalid date format: {date_str}")
    raise ValueError(
        f"Invalid date format: {date_str}. Expected YYYY-MM-DD, DD/MM/YYYY, or DD-MM-YYYY"
    )

# Helper function: Calculate years of term
def calculate_years_of_term(start_date):
    """
    Calculate years of term dynamically based on any start date provided.
    For the fifth year, add one extra month to the end date.

    Args:
        start_date (str): Start date in YYYY-MM-DD format.

    Returns:
        list: List of tuples containing (start_date, end_date) for each year.
    """
    start_obj = datetime.strptime(start_date, "%Y-%m-%d")
    dates = []

    current_date = start_obj
    for year in range(7):
        # Start date for the current year
        year_start = current_date

        # End date is one year minus one day from the start date
        year_end = year_start + relativedelta(years=1) - timedelta(days=1)

        # Add an extra month to the fifth subsequent ending date
        if year == 4:  # The fifth year (0-indexed)
            year_end += relativedelta(months=1)

        # Format and append the dates
        dates.append(
            (year_start.strftime("%d/%m/%Y"), year_end.strftime("%d/%m/%Y"))
        )

        # Move to the next year's start date
        current_date = year_end + timedelta(days=1)

    return dates

# Helper function: Calculate remainder dates
def calculate_remainder(years_of_term):
    if len(years_of_term) < 7:
        raise ValueError(
            "Not enough terms to calculate the fifth subsequent ending date.")

    fifth_end_date = datetime.strptime(years_of_term[4][1], "%d/%m/%Y")
    start_of_month = fifth_end_date.replace(day=1)
    next_month = start_of_month + timedelta(days=31)
    end_of_month = next_month.replace(day=1) - timedelta(days=1)
    return start_of_month.strftime("%d/%m/%Y"), end_of_month.strftime(
        "%d/%m/%Y")

# Serve the frontend page
@app.route("/")
def index():
    return render_template("frontend.html")

@app.route("/generate", methods=["POST"])
def generate_lease():
    try:

        # Collect data from the UI
        data = request.json
        tenant_name = data.get("tenant_name").upper()  # Convert to uppercase
        phone_number = data.get("phone_number")
        email_address = data.get("email_address")
        physical_address = data.get("physical_location")
        date_of_lease_entry = data.get("date_of_lease_entry", None)
        start_date = data.get("start_date", None)
        end_date = data.get("end_date", None)
        floor_plan = data.get("floor_plan")
        office_number = data.get("office_number")
        floor_number = data.get("floor_number").upper()
        po_box = data.get("po_box")
        post_code = data.get("post_code")
        town = data.get("town")
        parking_capacity = data.get("parking_capacity")
        escalation_rate = (data.get("escalation_rate"))
        escalation_type = data.get("type_of_escalation")
        yearly_rent = int(data.get("yearly_rent", 0))
        monthly_rent = yearly_rent // 12
        new_or_renew = data.get("new_or_renew")

        # Validate input
        if not start_date or not end_date:
            return jsonify({"error":
                            "Start date and end date are required."}), 400

        # Parse and format dates
        start_date_obj, start_date_formatted = parse_and_format_date(
            start_date)
        end_date_obj, end_date_formatted = parse_and_format_date(end_date)

        # Calculate years of term
        years_of_term = calculate_years_of_term(start_date)
        if len(years_of_term) < 7:
            return jsonify({
                "error":
                "Insufficient terms to calculate the fifth subsequent ending date."
            }), 400

        # Calculate fifth end date and remainder dates
        fifth_end_date = datetime.strptime(years_of_term[4][1], "%d/%m/%Y")
        duration = fifth_end_date - start_date_obj
        total_days = duration.days
        months = total_days // 30
        remaining_days = total_days % 30
        lease_duration = f"{months} months {remaining_days} days"
        remainder_dates = calculate_remainder(years_of_term)

        # Standardize date formats in years_of_term
        formatted_years = []
        for start, end in years_of_term:
            start_date = start
            end_date = end
            formatted_years.append((start_date, end_date))

        # Parse escalation
        def parse_escalation_rate(rate_str):
            """
            Parse escalation rate from a string like 'Ten(10)' to a float (e.g., 0.10).
            """
            match = re.search(
                r'\((\d+)\)',
                rate_str)  # Extract the numeric portion in parentheses
            if match:
                rate_percentage = int(match.group(1))
                return rate_percentage / 100  # Convert percentage to decimal
            raise ValueError(f"Invalid rate format: {rate_str}")

        # Calculate rent escalation
        def calculate_escalation(base_rent,
                                 rate_str,
                                 escalation_type,
                                 terms=7):
            try:
                # Parse escalation rate
                rate = parse_escalation_rate(rate_str)
                base_rent = int(base_rent)

                # Normalize escalation type
                normalized_type = escalation_type.strip().lower().replace(
                    "-", "").replace(" ", "")

                # Ensure valid escalation type
                valid_types = ["yearly", "afterfirsttwoyears", "everytwoyears"]
                if normalized_type not in valid_types:
                    logger.error(
                        f"Unknown escalation type received: {escalation_type}")
                    raise ValueError(
                        f"Invalid escalation type: {escalation_type}")

                # Initialize escalation logic
                initial_rent = base_rent
                escalations = [initial_rent]
                for year in range(1, terms):
                    if normalized_type == "yearly":
                        current_rent = initial_rent + (initial_rent * rate * year)
                    elif normalized_type == "afterfirsttwoyears" and year >= 2:
                        current_rent = initial_rent + (initial_rent * rate * (year - 1))
                    elif normalized_type == "everytwoyears" and year % 2 == 0:
                        current_rent = initial_rent + (initial_rent * rate * (year // 2))
                    else:
                        current_rent = initial_rent
                    escalations.append(int(current_rent))

                logger.info(
                    f"Calculated Escalations for {escalation_type}: {escalations}"
                )
                return escalations

            except Exception as e:
                logger.error(f"Error calculating escalation: {e}",
                             exc_info=True)
                return [base_rent] * terms

        escalated_rents = calculate_escalation(yearly_rent, escalation_rate,
                                               escalation_type)

        # Convert dates to words
        def date_to_words(date_str):
            """Convert date to words, handling multiple input formats"""
            date_obj, _ = parse_and_format_date(date_str)

            day = int(date_obj.strftime("%d"))
            month = date_obj.strftime("%B")
            year = int(date_obj.strftime("%Y"))

            day_words = {
                1: "First",
                2: "Second",
                3: "Third",
                4: "Fourth",
                5: "Fifth",
                6: "Sixth",
                7: "Seventh",
                8: "Eighth",
                9: "Ninth",
                10: "Tenth",
                11: "Eleventh",
                12: "Twelfth",
                13: "Thirteenth",
                14: "Fourteenth",
                15: "Fifteenth",
                16: "Sixteenth",
                17: "Seventeenth",
                18: "Eighteenth",
                19: "Nineteenth",
                20: "Twentieth",
                21: "Twenty-first",
                22: "Twenty-second",
                23: "Twenty-third",
                24: "Twenty-fourth",
                25: "Twenty-fifth",
                26: "Twenty-sixth",
                27: "Twenty-seventh",
                28: "Twenty-eighth",
                29: "Twenty-ninth",
                30: "Thirtieth",
                31: "Thirty-first"
            }.get(day)

            return f"{day_words} of {month} {num2words(year, lang='en').capitalize()}"

        def format_date(date_str):
            """Format date string to DD/MM/YYYY format"""
            if not date_str:
                return ""
            _, formatted_date = parse_and_format_date(date_str)
            return formatted_date

        start_date_words = date_to_words(start_date)
        fifth_end_date = years_of_term[4][1]
        end_date_words = date_to_words(fifth_end_date)

        # Convert numbers to words
        def number_to_words(number):
            from num2words import num2words
            return num2words(number, lang="en").capitalize()

        yearly_rent_words = number_to_words(yearly_rent)
        monthly_rent_words = number_to_words(monthly_rent)

        yearly_rent_text = f"{yearly_rent_words}: KSH {yearly_rent}"
        monthly_rent_text = f"{monthly_rent_words}: KSH {monthly_rent}"

        # Ensure the template exists
        script_dir = os.path.dirname(os.path.abspath(__file__))
        template_path = os.path.join(script_dir,
                                     "Golden Ivy Lease Template.docx")
        output_path = os.path.join(script_dir, "generated_lease.docx")

        if not os.path.exists(template_path):
            return jsonify({"error": "Template file not found!"}), 500

        # Load the lease document template
        document = Document(template_path)

        # Replacing escalation terms
        def replace_escalation_terms(escalated_rents):
            """
            Generate replacements for escalation terms in the document template.
            """
            escalation_replacements = {}

            def number_to_words(number):
                """Convert a number to words."""
                return num2words(number, lang="en").capitalize()

            def get_ordinal_suffix(number):
                """Get ordinal suffix for a number."""
                if 11 <= (number % 100) <= 13:
                    suffix = "th"
                else:
                    suffix = {1: "st", 2: "nd", 3: "rd"}.get(number % 10, "th")
                return f"{number}{suffix}"

            for i, rent in enumerate(escalated_rents[1:],
                                     start=2):  # Start from the 2nd year
                ordinal_suffix = get_ordinal_suffix(i)
                rent_words = number_to_words(rent)
                monthly_rent = rent // 12
                monthly_rent_words = number_to_words(monthly_rent)

                yearly_key = f"{ordinal_suffix} Year of Term Yearly Calculation"
                monthly_key = f"{ordinal_suffix} Year of Term Monthly Calculation"

                escalation_replacements[
                    yearly_key] = f"{rent_words} Only: KSH ({rent}/-)"
                escalation_replacements[
                    monthly_key] = f"KSH {monthly_rent}/- Monthly Rent"
                # Add new replacement for remaining year monthly calculation
                final_monthly_rent = escalated_rents[-1] // 12  # Get monthly rent from last year
                final_monthly_words = monthly_rent_words
                escalation_replacements["Remaining Year of Term Monthly Calculation"] = f"{final_monthly_words} Only (KSH {final_monthly_rent}/-)"

            return escalation_replacements

        # Formatted functions for floor number, rate of escalation and type of escalation, lease beginning date
        def format_lease_beginning_date(date_str):
            """Format date as '1st January 2025'"""
            date_obj, _ = parse_and_format_date(date_str)
            day = date_obj.day
            suffix = {
                1: 'st',
                2: 'nd',
                3: 'rd'
            }.get(day % 10 if day % 100 not in [11, 12, 13] else 0, 'th')
            return f"{day}{suffix} {date_obj.strftime('%B %Y')}"

        def format_floor_number(floor_num):
            """Convert floor number to ordinal format"""
            number_mapping = {
                'one': 'First',
                'two': 'Second',
                'three': 'Third',
                'four': 'Fourth',
                'five': 'Fifth',
                'six': 'Sixth',
                'seven': 'Seventh',
                'eight': 'Eighth',
                'nine': 'Ninth',
                'ten': 'Tenth',
                'eleven': 'Eleventh',
                'twelve': 'Twelfth',
                'thirteen': 'Thirteenth',
                'fourteen': 'Fourteenth',
                'fifteen': 'Fifteenth',
                'sixteen': 'Sixteenth',
                'seventeen': 'Seventeenth',
                'eighteen': 'Eighteenth',
                'nineteen': 'Nineteenth',
                'twenty': 'Twentieth',
                'twenty one': 'Twenty-first',
                'twenty two': 'Twenty-second',
                'twenty three': 'Twenty-third',
                'twenty four': 'Twenty-fourth',
                'twenty five': 'Twenty-fifth',
                'twenty six': 'Twenty-sixth',
                'twenty seven': 'Twenty-seventh',
                'twenty eight': 'Twenty-eighth',
                'twenty nine': 'Twenty-ninth',
                'thirty': 'Thirtieth',
                'attic': 'Attic',
                'Attic': 'Attic',
                'ground': 'Ground',
                'Ground': 'Ground',
                'basement': 'Basement',
                'Basement': 'Basement',
            }
            floor = floor_num.lower().replace('floor', '').strip()
            return f"{number_mapping.get(floor, floor)} Floor"

        def format_rate_of_escalation(rate_str):
            """Convert rate format from 'Ten(10)%' to 'ten(10)'"""
            match = re.search(r'(\w+)\((\d+)\)%', rate_str)
            if match:
                word, number = match.groups()
                return f"{word.lower()}({number})"
            return rate_str

        def format_type_of_escalation(escalation_type):
            """Format type of escalation according to specified rules"""
            type_mapping = {
                "yearly": "first year",
                "afterfirsttwoyears": "after the first two years",
                "everytwoyears": "every two years"
            }
            return type_mapping.get(escalation_type.lower().replace(" ", ""),
                                    escalation_type)

        # Generate replacements
        def generate_replacements(data, escalated_rents):
            # Calculate years of term
            years_of_term = calculate_years_of_term(data.get("start_date"))

            # Calculate remainder dates
            remainder_dates = calculate_remainder(years_of_term)

            # Get fifth end date
            fifth_end_date = years_of_term[4][1] if len(
                years_of_term) >= 5 else None
            """
            Generate a comprehensive replacements' dictionary.

            Args:
                data (dict): Input data from the frontend
                escalated_rents (list): List of calculated rents for each year of the term

            Returns:
                dict: Comprehensive dictionary of replacements
            """
            # Basic replacements from input data
            replacements = {

        # Rest of your existing replacements
        "Tenant Name": data.get("tenant_name", "").upper(),
        "Phone Number": data.get("phone_number", ""),
        "Email Address": data.get("email_address", ""),
        "Physical Location": data.get("physical_location", ""),
        "Office Number Page 1": data.get("office_number", "").upper(),
        "Office Number Page 5": data.get("office_number", ""),
        "Office Number": data.get("office_number", ""),
        "Floor Number Page 1": data.get("floor_number", "").upper(),
        "Floor Number": data.get("floor_number", ""),
        "Date of Lease Entry": format_lease_beginning_date(data["date_of_lease_entry"]),
        "Persons Names": data.get("tenant_name", "").upper(),
        "Start Date": format_date(data.get("start_date")),
        "End Date": format_date(data.get("fifth_end_date")),
        "Start_Date_in_words":date_to_words(data.get('start_date')),
        "End_Date_in_words":date_to_words(fifth_end_date) if fifth_end_date else '',
        "Starts Dates": format_date(data.get("start_date")),
        "New or Renew": data.get("new_or_renew", ""),
        "Yearly Rent": f"{number_to_words(yearly_rent)} Only: KSH ({yearly_rent}/-)",
        "Months Rent": f"KSH {monthly_rent}/- Monthly Rent",
        "Lease Term": data.get("lease_duration", ""),
        "PO Box number": data.get("po_box", ""),
        "post code": data.get("post_code", ""),
        "Town of residence": data.get("town", ""),
        "Floor plan in Sq foot": data.get("floor_plan", ""),
        "Parking Capacity": data.get("parking_capacity", ""),
        "Parking Capacity Page 5": data.get("parking_capacity", ""),
        "Rate of escalation": format_rate_of_escalation(data["escalation_rate"]),
        "Type of escalation": format_type_of_escalation(data["type_of_escalation"]),
        "Floor of office": format_floor_number(data["floor_number"]),
        "Lease Beginning Date": format_lease_beginning_date(data["start_date"]),
        "First Subsequent Ending Date": formatted_years[0][1],
        "Second Subsequent Starting Date": formatted_years[1][0],
        "Second Subsequent Ending Date": formatted_years[1][1],
        "Third Subsequent Starting Date": formatted_years[2][0],
        "Third Subsequent Ending Date": formatted_years[2][1],
        "Fourth Subsequent Starting Date": formatted_years[3][0],
        "Fourth Subsequent Ending Date": formatted_years[3][1],
        "Fifth Subsequent Starting Date": formatted_years[4][0],
        "Fifth Subsequent Ending Date": formatted_years[4][1],
        "Sixth Subsequent Starting Date": formatted_years[5][0],
        "Sixth Subsequent Ending Date": formatted_years[5][1],
        "Remainder Beginning Date": remainder_dates[0],
        "Remainder Ending Date": remainder_dates[1],
        # Terms to be underlined
        "Fifth Subsequent Finishing Date": formatted_years[4][1],
        "1st Years of Terms:": f"1st Year of Term: ({formatted_years[0][0]} to {formatted_years[0][1]}):",
        "2nd Years of Terms:": f"2nd Year of Term: ({formatted_years[1][0]} to {formatted_years[1][1]}):",
        "3rd Years of Terms:": f"3rd Year of Term: ({formatted_years[2][0]} to {formatted_years[2][1]}):",
        "4th Years of Terms:": f"4th Year of Term: ({formatted_years[3][0]} to {formatted_years[3][1]}):",
        "5th Years of Terms:": f"5th Year of Term: ({formatted_years[4][0]} to {formatted_years[4][1]}):",
        "One (1) Month being the remainder of the term": f"One (1) Month being the remainder of the term: ({remainder_dates[0]} to {remainder_dates[1]}):",
        # Terms to be made bold
        "LETTING OF OFFICE": "LETTING OF OFFICE",
        "designated Office": "designated Office",
        "designated parking spaces": "designated parking spaces",

    }

            # Add escalation-related replacements
            escalation_replacements = replace_escalation_terms(escalated_rents)

            # Merge the dictionaries
            replacements.update(escalation_replacements)

            return replacements

        # In the generate_lease function, replace the existing replacements creation with:
        replacements = generate_replacements(data, escalated_rents)


        # Replace Text in document and add style formatting
        def replace_text_with_formatting(document, replacements):
            # Define formatting rules: keys and their corresponding styles
            formatting_map = {
                "LETTING OF OFFICE": {'bold': True},
                "Office Number Page 1": {'bold': True},
                "Floor Number Page 1": {'bold': True},
                "Persons Names": {'bold': True},
                "Office Number Page 5": {'bold': True},
                "designated Office": {'bold': True},
                "Parking Capacity Page 5": {'bold': True},
                "designated parking spaces": {'bold': True},
                "Start_Date_in_words": {'underline': WD_UNDERLINE.SINGLE},
                "End_Date_in_words": {'underline': WD_UNDERLINE.SINGLE},
                "Starts Dates": {'underline': WD_UNDERLINE.SINGLE},
                "Fifth Subsequent Finishing Date": {'underline': WD_UNDERLINE.SINGLE},
                "1st Years of Terms:": {'underline': WD_UNDERLINE.SINGLE},
                "2nd Years of Terms:": {'underline': WD_UNDERLINE.SINGLE},
                "3rd Years of Terms:": {'underline': WD_UNDERLINE.SINGLE},
                "4th Years of Terms:": {'underline': WD_UNDERLINE.SINGLE},
                "5th Years of Terms:": {'underline': WD_UNDERLINE.SINGLE},
                "One (1) Month being the remainder of the term": {'underline': WD_UNDERLINE.SINGLE}
            }
            formatting_keys = list(formatting_map.keys())

            def replace_in_paragraph(paragraph, is_table=False):
                original_text = paragraph.text  

                # Process Tenant Name with bold
                if "Tenant Name" in original_text:
                    tenant_name = replacements.get("Tenant Name", "Tenant Name")
                    new_text = original_text.replace("Tenant Name", tenant_name)
                    parts = new_text.split(tenant_name)
                    paragraph.clear()
                    for i, part in enumerate(parts):
                        if part:
                            run = paragraph.add_run(part)
                            run.bold = False
                            run.font.size = Pt(12 if is_table else 14)
                        if i < len(parts) - 1:
                            run = paragraph.add_run(tenant_name)
                            run.bold = True
                            run.font.size = Pt(12 if is_table else 14)
                    original_text = paragraph.text 
                            
                # Process formatting keys using regex
                if any(key in original_text for key in formatting_keys):
                    pattern = re.compile(r'(' + '|'.join(map(re.escape, formatting_keys)) + r')')
                    parts = pattern.split(original_text)
                    paragraph.clear()
                    for part in parts:
                        if not part:
                            continue
                        if part in formatting_keys:
                            value = replacements.get(part, part)
                            run = paragraph.add_run(str(value))
                            style = formatting_map[part]
                            run.bold = style.get('bold', False)
                            if 'underline' in style:
                                run.underline = style['underline']
                        else:
                            paragraph.add_run(part)
                    original_text = paragraph.text

                # Handle non-formatting replacements differently for tables
                if is_table:
                    # Combine all runs to handle split placeholders
                    combined_text = ''.join(run.text for run in paragraph.runs)
                    new_text = combined_text
                    for key, value in replacements.items():
                        if key not in formatting_keys and key != "Tenant Name":
                            new_text = new_text.replace(key, str(value))
                    if new_text != combined_text:
                        paragraph.clear()
                        for text in new_text.split('\n'):
                            run = paragraph.add_run(text)
                            run.font.size = Pt(12)
                else:
                    # Existing per-run replacement for non-table text
                    for run in paragraph.runs:
                        current_text = run.text
                        new_text = current_text
                        for key, value in replacements.items():
                            if key not in formatting_keys and key != "Tenant Name":
                                new_text = new_text.replace(key, str(value))
                        if new_text != current_text:
                            run.text = new_text

            # Apply replacements to all document elements
            for paragraph in document.paragraphs:
                replace_in_paragraph(paragraph)

            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_in_paragraph(paragraph, is_table=True)

            for section in document.sections:
                for header_footer in [section.header, section.footer]:
                    if header_footer:
                        for paragraph in header_footer.paragraphs:
                            replace_in_paragraph(paragraph)
        replace_text_with_formatting(document, replacements)
        
        # logging of keys that have not been replaced
        def log_unmatched_keys(replacements, document):
            """
            Log placeholders that are not matched in the document.
            """
            original_text = ""

            # Collect all text from the document
            for paragraph in document.paragraphs:
                original_text += paragraph.text
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            original_text += paragraph.text
            for section in document.sections:
                for header_footer in [
                        section.header, section.footer,
                        section.first_page_header, section.first_page_footer
                ]:
                    if header_footer:
                        for paragraph in header_footer.paragraphs:
                            original_text += paragraph.text

            unmatched_keys = [
                key for key in replacements if key not in original_text
            ]
            if unmatched_keys:
                logger.warning(
                    f"Unmatched placeholders in the document: {unmatched_keys}"
                )

        # Load the lease document template
        document = Document(template_path)

        # Call the function to replace placeholders in the document
        replace_text_with_formatting(document, replacements)

        # Log any unreplaced keys - Add this line right before saving the document
        log_unmatched_keys(replacements, document)

        # Save the updated document
        document.save(output_path)

        # Right before returning the file, create a JSON response with the calculated values
        response = send_file(output_path, as_attachment=True)
        response.headers['X-Fifth-End-Date'] = years_of_term[4][1]  # Fifth subsequent ending date
        response.headers['X-Lease-Duration'] = lease_duration
        return response

    except Exception as e:
        logger.error(f"Lease generation error: {str(e)}", exc_info=True)
        return jsonify({
            "error": "Failed to generate the lease. Please check your input and try again."
        }), 500

@app.route("/calculate-dates", methods=["POST"])
def calculate_dates():
    try:
        start_date = request.json.get("start_date")
        if not start_date:
            return jsonify({"error": "Start date is required"}), 400

        years_of_term = calculate_years_of_term(start_date)
        start_date_obj = datetime.strptime(start_date, "%Y-%m-%d")
        fifth_end_date = datetime.strptime(years_of_term[4][1], "%d/%m/%Y")
        
        # Calculate exact duration
        duration = fifth_end_date - start_date_obj
        total_days = duration.days
        months = total_days // 30
        remaining_days = total_days % 30

        lease_duration = f"{months} months {remaining_days} days"

        return jsonify({
            "fifthEndDate": years_of_term[4][1],
            "leaseDuration": lease_duration
        })

    except Exception as e:
        logger.error(f"Date calculation error: {str(e)}", exc_info=True)
        return jsonify({"error": str(e)}), 500

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=8000)




